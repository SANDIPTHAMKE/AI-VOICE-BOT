from flask import Flask, request, render_template
import speech_recognition as sr
import time
import pyttsx3
import datetime
import webbrowser
import os
from openpyxl import Workbook
from docx import Document

app = Flask(__name__)

# Initialize the recognizer and the text-to-speech engine
recognizer = sr.Recognizer()
engine = pyttsx3.init()

def speak(text):
    engine.say(text)
    engine.runAndWait()

#Home page of web application
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/run-python', methods=['POST'])
def run_python():
    wish_me()
    run_once = 1
    query = ""
    while run_once == 1:
        run_once += 1
        query = listen()
        if query in ["exit", "stop"]:
            speak("Thank you. You are a good speaker. Goodbye! Have a nice time.")
            break
        perform_task(query)
    return render_template('index.html', command=query)

def listen(slowdown_factor=1):
    with sr.Microphone() as source:
        print("Listening...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)

        try:
            time.sleep(slowdown_factor)
            print("Recognizing...")
            query = recognizer.recognize_google(audio, language='en-US')
            print(f"User said: {query}")
            return query.lower()
        
        except sr.UnknownValueError:
            speak("Sorry, I did not catch that. Could you please repeat?")
        except sr.RequestError:
            speak("Sorry, the service is down. Please try again later.")
        return ""

def wish_me():
    hour = datetime.datetime.now().hour
    greeting = "Good Morning!" if hour < 12 else "Good Afternoon!" if hour < 18 else "Good Evening!"
    speak(greeting)
    speak("I am your personal assistant. How can I help you today?")

def perform_task(query):
    if "time" in query:
        current_time = datetime.datetime.now().strftime("%I:%M %p")
        speak(f"The time is {current_time}")
    #Open Notepad and write what the user says and update it in text document.
    elif "open notepad" in query:
        speak("Opening Notepad. Start dictating, and I will write for you. Say stop writing when you are done.")
        write_to_notepad()
    #Open word and write what the user says and update it in word document.
    elif "open word" in query:
        speak("Opening Microsoft Word. Start dictating, and I will write for you. Say stop writing when you are done.")
        create_word_document()

    elif "open excel" in query:
        speak("Opening Microsoft Excel.")
        write_to_excel()
    elif "open powerpoint" in query:
        speak("Opening Microsoft PowerPoint.")
        os.system("start powerpnt")
    elif "open vscode" in query or "open visual studio code" in query:
        speak("Opening Visual Studio Code.")
        os.system("code")
    elif "open whatsapp" in query:
        speak("Opening WhatsApp.")
        os.system("WhatsApp")
    elif "search for" in query:
        search_term = query.replace("search for", "").strip()
        speak(f"Searching for {search_term} on Google")
        webbrowser.open(f"https://www.google.com/search?q={search_term.replace(' ', '+')}")
    elif "play" in query and "youtube" in query:
        song_name = query.replace("play", "").replace("on youtube", "").strip()
        speak(f"Playing {song_name} on YouTube")
        webbrowser.open(f"https://www.youtube.com/results?search_query={song_name.replace(' ', '+')}")
    elif "open" in query:
        website = query.replace("open", "").strip().replace(" ", "")
        speak(f"Opening {website}")
        webbrowser.open(f"https://{website}.com")
    elif "say" in query:
        speak(query.replace("say", "").strip())
    else:
        speak("Sorry, I can not help with that right now.")

def speak(text):
    os.system(f'''powershell -c "Add-Type -AssemblyName System.Speech; $speak = New-Object System.Speech.Synthesis.SpeechSynthesizer;$speak.SelectVoice('Microsoft Zira Desktop');  $speak.Speak(\'{text}\');"''')

def write_to_notepad():
    # Open Notepad
    os.system("notepad.exe")  # Wait for Notepad to open

    while True:
        query = listen()
        if "stop writing" in query:
            speak("Stopped writing.")
            break
        else:
            # Send the recognized text to Notepad
            os.system(f'''powershell -c "$wshell = New-Object -ComObject wscript.shell; $wshell.SendKeys('{query} ')"''')
def create_word_document():
    doc = Document()
    paragraph = doc.add_paragraph()
    
    while True:
        query = listen()  # Assuming this is your voice recognition function
        if "stop writing" in query:
            speak("Stopped writing.")
            break
        else:
            paragraph.add_run(query + " ")
            doc.save('output.docx')  # Save the document after every addition to avoid losing data
    else:
        doc = Document()
        paragraph = doc.add_paragraph()

        while True:
            query = listen()
            if "stop writing" in query:
                speak("Stopped writing.")
                break
            else:
                paragraph.add_run(query + " ")
                doc.save('example.docx')


def write_to_excel():
    # Create a new workbook and select the active worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "User Data"

    # Ask for the data to write in Excel
    speak("What data would you like to enter into Excel?")
    data = listen()

    # Writing each sentence or data item into a separate row
    rows = data.split(",")  # Assuming user inputs comma-separated data for multiple rows
    for index, row in enumerate(rows, start=1):
        # Split each row into columns (assuming space-separated values)
        columns = row.strip().split(" ")
        for col_index, value in enumerate(columns, start=1):
            ws.cell(row=index, column=col_index).value = value

    # Save the workbook
    filename = "user_data.xlsx"
    wb.save(filename)
    speak(f"Data has been saved in {filename}")

    # Open the Excel file automatically
    os.system(f'start excel "{filename}"')
def write_data(app_name):
    data = listen()
    if app_name == "notepad":
        with open("output.txt", "w") as file:
            file.write(data)
        os.system("notepad output.txt")
    else:
        speak("Currently, writing to Word, Excel, or PowerPoint is not automated.")

if __name__ == "__main__":
    app.run(debug=True)
